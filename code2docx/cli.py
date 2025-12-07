#!/usr/bin/env python3
"""
code2docx -- Darcula-themed code export to DOCX

- Filenames: Times New Roman, 14pt bold
- Code: token-colored via Pygments -> inserted as runs (Courier New, 10pt)
- Paragraph background: Darcula dark (#2b2b2b)
- After files: "Screenshots:" with images in 2-column layout, then "Done."
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pygments import lex
from pygments.util import ClassNotFound
from pygments.lexers import guess_lexer_for_filename, TextLexer

OUTPUT_NAME = "project_files.docx"

# Image extensions to look for
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp'}

# Darcula-inspired token color map (hex strings)
TOKEN_COLOR_MAP = {
    "Comment": "808080",
    "Keyword": "cc7832",
    "Name.Function": "ffc66d",
    "Name.Class": "ffc66d",
    "Name.Builtin": "a9b7c6",
    "String": "6a8759",
    "Number": "6897bb",
    "Operator": "a9b7c6",
    "Punctuation": "a9b7c6",
    "Name": "a9b7c6",
    "Text": "a9b7c6",
    "Error": "ff0000",
    "Literal": "6a8759",
}

CODE_BG_HEX = "2b2b2b"
IMAGE_WIDTH_INCHES = 3.0


def hex_to_rgb(hexstr):
    hexstr = hexstr.strip().lstrip("#")
    r = int(hexstr[0:2], 16)
    g = int(hexstr[2:4], 16)
    b = int(hexstr[4:6], 16)
    return RGBColor(r, g, b)


def token_color(tok_type):
    s = str(tok_type)
    for key in ("Name.Function", "Name.Class", "Name.Builtin"):
        if key in s and key in TOKEN_COLOR_MAP:
            return hex_to_rgb(TOKEN_COLOR_MAP[key])
    for key in TOKEN_COLOR_MAP:
        if key in s:
            return hex_to_rgb(TOKEN_COLOR_MAP[key])
    return hex_to_rgb(TOKEN_COLOR_MAP["Text"])


def set_paragraph_shading(paragraph, fill_color):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    existing = pPr.find(qn('w:shd'))
    if existing is not None:
        pPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def add_filename(doc, filename):
    p = doc.add_paragraph()
    run = p.add_run(filename)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(14)
    run.font.bold = True


def add_colored_code_block(doc, code_text, filename_hint):
    try:
        lexer = guess_lexer_for_filename(filename_hint, code_text)
    except (ClassNotFound, Exception):
        lexer = TextLexer()

    tokens = list(lex(code_text, lexer))
    lines = [[]]
    
    for tok_type, value in tokens:
        value = value.replace("\t", " " * 4)
        parts = value.splitlines(keepends=True)
        for part in parts:
            if part.endswith("\n") or part.endswith("\r\n"):
                text_part = part.rstrip("\r\n")
                if text_part:
                    lines[-1].append((tok_type, text_part))
                lines.append([])
            else:
                if part:
                    lines[-1].append((tok_type, part))
    
    if len(lines) > 1 and lines[-1] == []:
        lines.pop()

    for line_tokens in lines or [[]]:
        p = doc.add_paragraph()
        set_paragraph_shading(p, CODE_BG_HEX)

        if not line_tokens:
            r = p.add_run(" ")
            r.font.name = "Courier New"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            r.font.size = Pt(10)
            r.font.color.rgb = hex_to_rgb(TOKEN_COLOR_MAP["Text"])
            continue

        for tok_type, txt in line_tokens:
            run = p.add_run(txt)
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            run.font.size = Pt(10)
            run.font.color.rgb = token_color(tok_type)


def is_text_file(path, blocksize=1024):
    try:
        with open(path, "rb") as f:
            chunk = f.read(blocksize)
            if b'\x00' in chunk:
                return False
    except Exception:
        return False
    return True


def is_image_file(filename):
    ext = os.path.splitext(filename)[1].lower()
    return ext in IMAGE_EXTENSIONS


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_screenshots_section(doc, cwd):
    image_files = []
    entries = sorted(os.listdir(cwd))
    
    for name in entries:
        if name.startswith('.'):
            continue
        if name.lower().endswith('.docx'):
            continue
        path = os.path.join(cwd, name)
        if os.path.isfile(path) and is_image_file(name):
            image_files.append((name, path))
    
    if not image_files:
        doc.add_paragraph("No screenshots found in folder.")
        return
    
    num_rows = (len(image_files) + 1) // 2
    table = doc.add_table(rows=num_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    
    for idx, (name, path) in enumerate(image_files):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = table.rows[row_idx].cells[col_idx]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            run = paragraph.add_run()
            run.add_picture(path, width=Inches(IMAGE_WIDTH_INCHES))
            caption_para = cell.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(name)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.name = "Times New Roman"
        except Exception as e:
            error_run = paragraph.add_run(f"[Could not load: {name}]\n{e}")
            error_run.font.size = Pt(9)
            error_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()


def main():
    # Use current working directory (where command is run)
    cwd = os.getcwd()
    
    # Optional: accept path as argument
    if len(sys.argv) > 1:
        cwd = os.path.abspath(sys.argv[1])
    
    if not os.path.isdir(cwd):
        print(f"Error: '{cwd}' is not a valid directory")
        sys.exit(1)
    
    print(f"Processing folder: {cwd}")
    doc = Document()

    entries = sorted(os.listdir(cwd))
    any_files = False

    for name in entries:
        if name.startswith('.'):
            continue
        if name.lower().endswith('.docx'):
            continue
        if is_image_file(name):
            continue

        path = os.path.join(cwd, name)
        if not os.path.isfile(path):
            continue

        any_files = True
        print(f"  Adding: {name}")
        add_filename(doc, name)

        if is_text_file(path):
            try:
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
            except Exception as e:
                content = f"<Could not read file: {e}>"
                doc.add_paragraph(content)
                continue

            try:
                add_colored_code_block(doc, content, name)
            except Exception:
                p = doc.add_paragraph()
                set_paragraph_shading(p, CODE_BG_HEX)
                run = p.add_run(content)
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb(TOKEN_COLOR_MAP["Text"])
        else:
            doc.add_paragraph("<Binary or non-text file - skipped>")

        doc.add_paragraph()

    if not any_files:
        doc.add_paragraph("No code files found in this folder.")

    # Screenshots section
    p = doc.add_paragraph()
    run = p.add_run("Screenshots:")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    add_screenshots_section(doc, cwd)
    doc.add_paragraph("Done.")

    out_path = os.path.join(cwd, OUTPUT_NAME)
    doc.save(out_path)
    print(f"\n✓ Saved: {out_path}")


if __name__ == "__main__":
    main()